"""Format writers for cleaned manuscripts."""

from __future__ import annotations

import html
import json
import mimetypes
import re
import uuid
import zipfile
from collections.abc import Iterable
from pathlib import Path

from gutenberg_cleaner.reports import CleanupReport


def write_outputs(
    markdown: str,
    *,
    output_base: Path,
    formats: Iterable[str],
    title: str | None,
    author: str | None,
    report: CleanupReport,
) -> dict[str, Path]:
    """Write selected output formats and return their paths by format name."""
    paths: dict[str, Path] = {}
    output_base.parent.mkdir(parents=True, exist_ok=True)

    for fmt in formats:
        normalized = fmt.lower().strip().lstrip(".")
        if normalized == "md":
            path = output_base.with_suffix(".md")
            path.write_text(markdown, encoding="utf-8")
        elif normalized == "json":
            path = output_base.with_suffix(".report.json")
            path.write_text(
                json.dumps(report.to_dict(), indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
        elif normalized == "html":
            path = output_base.with_suffix(".html")
            path.write_text(markdown_to_html_document(markdown, title=title), encoding="utf-8")
        elif normalized == "epub":
            path = output_base.with_suffix(".epub")
            write_epub(markdown, path=path, title=title or "Untitled", author=author or "Unknown")
        elif normalized == "docx":
            path = output_base.with_suffix(".docx")
            write_docx(markdown, path=path, title=title, author=author)
        else:
            raise ValueError(f"Unsupported output format: {fmt}")
        paths[normalized] = path

    return paths


def markdown_to_html_document(markdown: str, *, title: str | None = None) -> str:
    body = "\n".join(_markdown_blocks_to_html(markdown))
    doc_title = html.escape(title or _first_heading(markdown) or "Cleaned manuscript")
    return (
        "<!doctype html>\n"
        "<html lang=\"en\">\n<head>\n"
        "  <meta charset=\"utf-8\">\n"
        f"  <title>{doc_title}</title>\n"
        "  <style>body{font-family:serif;line-height:1.55;max-width:42rem;"
        "margin:3rem auto;padding:0 1rem}"
        "h1,h2,h3{text-align:center} "
        ".scene-break{text-align:center;letter-spacing:.4em}</style>\n"
        "</head>\n<body>\n"
        f"{body}\n"
        "</body>\n</html>\n"
    )


def write_epub(markdown: str, *, path: Path, title: str, author: str) -> None:
    """Write a minimal EPUB 3 file with stdlib-only dependencies."""
    uid = f"urn:uuid:{uuid.uuid4()}"
    xhtml = markdown_to_html_document(markdown, title=title).replace("<!doctype html>\n", "")
    xhtml = xhtml.replace(
        "<html lang=\"en\">",
        '<html xmlns="http://www.w3.org/1999/xhtml" lang="en">',
    )
    nav = _build_nav(markdown)
    opf = _build_opf(title=title, author=author, uid=uid)

    with zipfile.ZipFile(path, "w") as epub:
        epub.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        epub.writestr("META-INF/container.xml", _container_xml())
        epub.writestr("EPUB/package.opf", opf)
        epub.writestr("EPUB/nav.xhtml", nav)
        epub.writestr("EPUB/text.xhtml", xhtml)


def write_docx(markdown: str, *, path: Path, title: str | None, author: str | None) -> None:
    """Write DOCX via optional python-docx dependency."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - exercised only without optional extra
        raise RuntimeError("DOCX output requires: pip install 'gutenberg-cleaner[docx]'") from exc

    document = Document()
    properties = document.core_properties
    if title:
        properties.title = title
    if author:
        properties.author = author

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            document.add_heading(
                _strip_markdown(heading.group(2)),
                level=min(len(heading.group(1)), 4),
            )
        elif stripped == "***":
            para = document.add_paragraph("***")
            para.alignment = 1
        else:
            document.add_paragraph(_strip_markdown(stripped))
    document.save(path)


def _markdown_blocks_to_html(markdown: str) -> Iterable[str]:
    paragraphs: list[str] = []
    heading_count = 0
    for block in re.split(r"\n\s*\n", markdown.strip()):
        block = block.strip()
        if not block:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", block)
        if heading:
            level = min(len(heading.group(1)), 6)
            heading_count += 1
            label = html.escape(_strip_markdown(heading.group(2)))
            yield f'<h{level} id="heading-{heading_count}">{label}</h{level}>'
        elif block == "***":
            yield '<p class="scene-break" aria-label="scene break">***</p>'
        else:
            safe = html.escape(_strip_markdown(block)).replace("\n", "<br>\n")
            paragraphs.append(f"<p>{safe}</p>")
            yield paragraphs.pop()


def _build_nav(markdown: str) -> str:
    items = []
    count = 0
    for line in markdown.splitlines():
        match = re.match(r"^(#{1,3})\s+(.+)$", line.strip())
        if match:
            count += 1
            label = html.escape(_strip_markdown(match.group(2)))
            items.append(f'<li><a href="text.xhtml#heading-{count}">{label}</a></li>')
    if not items:
        items.append('<li><a href="text.xhtml">Start</a></li>')
    return (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" lang="en">\n'
        '<head><title>Table of Contents</title></head>\n'
        '<body><nav epub:type="toc" xmlns:epub="http://www.idpf.org/2007/ops"><h1>Contents</h1><ol>'
        + "".join(items)
        + "</ol></nav></body></html>\n"
    )


def _build_opf(*, title: str, author: str, uid: str) -> str:
    return f'''<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="book-id">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="book-id">{html.escape(uid)}</dc:identifier>
    <dc:title>{html.escape(title)}</dc:title>
    <dc:creator>{html.escape(author)}</dc:creator>
    <dc:language>en</dc:language>
    <meta property="dcterms:modified">2026-01-01T00:00:00Z</meta>
  </metadata>
  <manifest>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="text" href="text.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine><itemref idref="text"/></spine>
</package>
'''


def _container_xml() -> str:
    return '''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="EPUB/package.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
'''


def _first_heading(markdown: str) -> str | None:
    match = re.search(r"^#\s+(.+)$", markdown, flags=re.MULTILINE)
    return _strip_markdown(match.group(1)) if match else None


def _strip_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text.strip()


mimetypes.add_type("application/epub+zip", ".epub")
